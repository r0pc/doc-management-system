"""Smoke tests for the corpus generator: files, formats, entity cross-check, determinism."""

import csv
import re
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from entities import SPECS
from generate_synthetic_corpus import generate_corpus
from templates import LABEL_PHRASES

CNIC_TEXT_RE = re.compile(r"(?<!\d)\d{5}-\d{7}-\d(?!\d)")
CARD_TEXT_RE = re.compile(r"(?<!\d)\d{16}(?!\d)")
ACCOUNT_TEXT_RE = re.compile(r"(?<!\d)PK[A-Z]{2}\d{13}(?!\d)")
PASSPORT_TEXT_RE = re.compile(r"(?<![A-Z0-9])[A-Z]{2}\d{7}(?!\d)")
# Literal "PKR" is reserved for salary lines across every template (see artifact_contract.md).
SALARY_TEXT_RE = re.compile(r"PKR\s\d[\d,]*")

COUNT_RES = {
    "cnic": CNIC_TEXT_RE,
    "card": CARD_TEXT_RE,
    "account": ACCOUNT_TEXT_RE,
    "passport": PASSPORT_TEXT_RE,
    "salary": SALARY_TEXT_RE,
}


def _generate(out: Path) -> Path:
    generate_corpus(count=6, out_dir=out, seed=42)
    return out / "manifest.json"


def _read_manifest_rows(manifest_csv: Path) -> list[dict[str, str]]:
    with manifest_csv.open(newline="", encoding="utf-8") as fh:
        return list(csv.DictReader(fh))


def test_manifest_and_three_formats_per_record(tmp_path: Path):
    sidecar_path = _generate(tmp_path)
    manifest_csv = tmp_path / "manifest.csv"
    assert manifest_csv.exists()
    assert sidecar_path.exists()

    rows = _read_manifest_rows(manifest_csv)
    assert len(rows) == 6
    assert set(rows[0]) == {
        "file_path",
        "label_doc_type",
        "label_level",
        "version_id",
        "generated_at",
    }

    for ext in ("docx", "xlsx", "pdf"):
        files = sorted(tmp_path.glob(f"*.{ext}"))
        assert len(files) == 6, f"expected 6 .{ext} files"


def test_all_three_formats_parse_back_non_empty_text(tmp_path: Path):
    _generate(tmp_path)
    for docx_path in sorted(tmp_path.glob("*.docx")):
        text = "\n".join(p.text for p in Document(docx_path).paragraphs)
        assert text.strip(), f"empty docx text: {docx_path.name}"
    for xlsx_path in sorted(tmp_path.glob("*.xlsx")):
        wb = load_workbook(xlsx_path)
        cells = [
            str(c) for row in wb.active.iter_rows(values_only=True) for c in row if c is not None
        ]
        assert any(c.strip() for c in cells), f"empty xlsx: {xlsx_path.name}"
    for pdf_path in sorted(tmp_path.glob("*.pdf")):
        assert pdf_path.read_bytes()[:5] == b"%PDF-"
    sidecar = __import__("json").loads((tmp_path / "manifest.json").read_text(encoding="utf-8"))
    for record in sidecar["records"]:
        assert record["source_text"].strip()


def test_strip_label_phrases_absent_from_bodies(tmp_path: Path):
    _generate(tmp_path)
    texts = ["\n".join(p.text for p in Document(p).paragraphs) for p in tmp_path.glob("*.docx")]
    sidecar = __import__("json").loads((tmp_path / "manifest.json").read_text(encoding="utf-8"))
    texts.extend(r["source_text"] for r in sidecar["records"])
    assert texts
    for text in texts:
        lowered = text.lower()
        for phrase in LABEL_PHRASES:
            assert phrase not in lowered, f"label phrase leaked: {phrase!r}"


def test_entity_counts_within_specs_and_cross_check_text(tmp_path: Path):
    _generate(tmp_path)
    sidecar = __import__("json").loads((tmp_path / "manifest.json").read_text(encoding="utf-8"))

    levels_seen: set[str] = set()
    for record in sidecar["records"]:
        level = record["level"]
        levels_seen.add(level)
        counts = record["entity_counts"]
        spec = SPECS[level]
        lo, hi = spec["cnic"]
        assert lo <= counts["cnic"] <= hi, f"cnic {counts['cnic']} outside {spec['cnic']} ({level})"
        lo, hi = spec["salary"]
        assert lo <= counts["salary"] <= hi
        lo, hi = spec["account"]
        assert lo <= counts["account"] <= hi
        expected_cards = (1, 2) if level == "restricted" else (0, 0)
        assert expected_cards[0] <= counts["card"] <= expected_cards[1]

        text = record["source_text"]
        for key, regex in COUNT_RES.items():
            found = len(regex.findall(text))
            assert found == counts[key], (
                f"{key}: text has {found}, sidecar says {counts[key]} (record {record['index']})"
            )
    assert levels_seen == set(SPECS)


def test_same_seed_produces_identical_manifest_bytes(tmp_path: Path):
    out_a, out_b = tmp_path / "a", tmp_path / "b"
    generate_corpus(count=6, out_dir=out_a, seed=42)
    generate_corpus(count=6, out_dir=out_b, seed=42)

    assert (out_a / "manifest.csv").read_bytes() == (out_b / "manifest.csv").read_bytes()
    assert (out_a / "manifest.json").read_bytes() == (out_b / "manifest.json").read_bytes()
