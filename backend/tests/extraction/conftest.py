"""Generated-in-memory fixtures for extraction tests — no binary fixtures committed.

Every fixture builds real bytes through the installed parsers (pymupdf,
python-docx, openpyxl, stdlib zipfile) so the suite stays hermetic.
"""

import io
import zipfile

import docx
import openpyxl
import pymupdf
import pytest


def _build_pdf(page_contents: list[str | None]) -> bytes:
    document = pymupdf.open()
    try:
        for content in page_contents:
            page = document.new_page()
            if content is not None:
                page.insert_text((72, 72), content)
        return document.tobytes()
    finally:
        document.close()


def _build_docx() -> bytes:
    document = docx.Document()
    document.add_paragraph("Invoice summary for vendor engagement.")
    document.add_paragraph("Payment terms follow the ledger schedule.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Qty"
    table.cell(1, 0).text = "Widgets"
    table.cell(1, 1).text = "12"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_xlsx(*, populated: bool) -> bytes:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Ledger"
    if populated:
        sheet["A1"] = "Total"
        sheet["B1"] = 42
        sheet["A2"] = "partial"  # B2 left empty -> None cell must be skipped
        notes = workbook.create_sheet("Notes")
        notes["A1"] = "audit note"
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


@pytest.fixture(scope="session")
def pdf_bytes() -> bytes:
    return _build_pdf(
        [
            "Quarterly report body text with plenty of characters.",
            "Second page continuation prose.",
        ]
    )


@pytest.fixture(scope="session")
def scanned_pdf_bytes() -> bytes:
    """A PDF whose pages carry no text layer — the OCR handoff case."""
    return _build_pdf([None])


@pytest.fixture(scope="session")
def docx_bytes() -> bytes:
    return _build_docx()


@pytest.fixture(scope="session")
def xlsx_bytes() -> bytes:
    return _build_xlsx(populated=True)


@pytest.fixture(scope="session")
def empty_xlsx_bytes() -> bytes:
    return _build_xlsx(populated=False)


@pytest.fixture(scope="session")
def plain_zip_bytes() -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("readme.txt", "just an archive")
    return buffer.getvalue()


@pytest.fixture(scope="session")
def custom_ooxml_zip_bytes() -> bytes:
    """Zip with an OOXML-style content-types part but no word/ or xl/ members."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("[Content_Types].xml", "<?xml version='1.0'?><Types/>")
        archive.writestr("custom/data.xml", "<data/>")
    return buffer.getvalue()


@pytest.fixture(scope="session")
def garbage_bytes() -> bytes:
    return b"\x00\x01\x02not-a-format\xff"
