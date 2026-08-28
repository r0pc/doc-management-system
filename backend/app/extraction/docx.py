"""DOCX text extraction via python-docx: paragraphs plus table cells.

A docx has no physical pages, so the whole document is reported as a single
pseudo-page numbered 1 (``PageText.page_no`` stays strictly ``int``).
"""

from __future__ import annotations

import io

from app.extraction.base import ExtractedDocument, PageText, ParserUnavailable
from app.extraction.sniff import MIME_DOCX


class DocxHandler:
    """Extracts paragraph text, then table cell text row by row."""

    def extract(self, data: bytes) -> ExtractedDocument:
        try:
            import docx
        except ImportError as exc:
            raise ParserUnavailable("python-docx is not installed") from exc
        document = docx.Document(io.BytesIO(data))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        text = "\n".join(parts)
        return ExtractedDocument(
            text=text,
            pages=[PageText(page_no=1, text=text)],
            mime_sniffed=MIME_DOCX,
            char_count=len(text),
            ocr_used=False,
        )
